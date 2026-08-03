from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\PRACTICAPROFESIONAL\outputs\Propuesta_Proyecto_SIGVITS.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT = "F4F6F9"
WHITE = "FFFFFF"
BLACK = "000000"


def set_font(run, name="Calibri", size=None, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            set_cell_margins(cell)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_font(r)
    return p


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{label}: ")
    set_font(r1, bold=True, color=INK)
    r2 = p.add_run(text)
    set_font(r2, color=INK)
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.333

for style_name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

bullet_style = styles["List Bullet"]
bullet_style.font.name = "Calibri"
bullet_style.font.size = Pt(11)
bullet_style.paragraph_format.left_indent = Inches(0.375)
bullet_style.paragraph_format.first_line_indent = Inches(-0.194)
bullet_style.paragraph_format.space_after = Pt(4)
bullet_style.paragraph_format.line_spacing = 1.208

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
hr = hp.add_run("SIGVITS | Propuesta académica de proyecto")
set_font(hr, size=9, color=MUTED, bold=True)

footer = section.footer
add_page_number(footer.paragraphs[0])

# Portada: proposal_centerpiece
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(24)
p.paragraph_format.space_after = Pt(8)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PROPUESTA DE PROYECTO")
set_font(r, size=12, color=MUTED, bold=True)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SIGVITS")
set_font(r, size=28, color=INK, bold=True)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Sistema de Gestión y Vigilancia de\nInfecciones de Transmisión Sexual")
set_font(r, size=17, color=BLUE, bold=True)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(24)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Piloto para el municipio de Puerto Cortés, Cortés")
set_font(r, size=11.5, color=MUTED, italic=True)

meta = doc.add_table(rows=3, cols=2)
set_table_geometry(meta, [4680, 4680])
remove_table_borders(meta)
metadata = [
    ("Estudiante", "____________________________", "Asesor académico", "____________________________"),
    ("Carrera", "____________________________", "Institución", "____________________________"),
    ("Lugar", "Puerto Cortés, Cortés", "Fecha", "29 de julio de 2026"),
]
for row, values in zip(meta.rows, metadata):
    for idx in range(2):
        cell = row.cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(7)
        label = values[idx * 2]
        value = values[idx * 2 + 1]
        r1 = p.add_run(f"{label}\n")
        set_font(r1, size=9, color=MUTED, bold=True)
        r2 = p.add_run(value)
        set_font(r2, size=10.5, color=BLACK)

doc.add_paragraph()
add_callout(
    doc,
    "Propósito",
    "Digitalizar y fortalecer el flujo institucional de información del programa ITS, desde la captura en los establecimientos de salud hasta la consolidación y el análisis para la toma de decisiones.",
)

doc.add_page_break()

doc.add_heading("1. Nombre del proyecto", level=1)
add_body(
    doc,
    "SIGVITS — Sistema de Gestión y Vigilancia de Infecciones de Transmisión Sexual.",
)
add_body(
    doc,
    "El proyecto consiste en el desarrollo de una plataforma web institucional orientada a la gestión de los formularios, reportes e indicadores del programa de Infecciones de Transmisión Sexual (ITS). La primera implementación se plantea como un piloto en Puerto Cortés, con cobertura para la coordinación municipal y sus 12 establecimientos de salud.",
)

doc.add_heading("2. Problema o necesidad que resolverá", level=1)
add_body(
    doc,
    "Actualmente, la captura, revisión y consolidación de la información del programa ITS depende en gran medida de formularios, archivos de Excel y comunicaciones realizadas por distintos medios. El mismo dato puede transcribirse varias veces a medida que pasa del establecimiento a la coordinación municipal, la región y el nivel central. Este proceso incrementa el tiempo de elaboración de informes y eleva el riesgo de errores, duplicidades y diferencias entre versiones.",
)
add_body(
    doc,
    "También existe una necesidad de trazabilidad: cuando un informe contiene inconsistencias, resulta difícil identificar quién realizó una modificación, cuál fue la observación o qué versión fue finalmente aprobada. A esto se suma la limitada disponibilidad de información oportuna para analizar tendencias, procedencias de pacientes, cobertura territorial e indicadores por período.",
)
add_body(doc, "Las principales necesidades identificadas son:")
for item in (
    "Reducir la digitación repetida y los errores durante la consolidación de los formularios ITS 1 e ITS 2.",
    "Estandarizar el flujo de envío, revisión, devolución, corrección, aprobación y cierre de los reportes.",
    "Proteger la información individual y sensible de los pacientes mediante accesos definidos por rol y territorio.",
    "Contar con reportes, indicadores, mapas y comparativos actualizados que apoyen la vigilancia epidemiológica y la toma de decisiones.",
    "Mantener un historial auditable de cambios, exportaciones y decisiones realizadas en cada nivel institucional.",
):
    add_bullet(doc, item)

doc.add_heading("3. Descripción general del sistema", level=1)
add_body(
    doc,
    "SIGVITS será una aplicación web que conectará el flujo de trabajo entre los establecimientos de salud, la coordinación municipal, la región sanitaria y el nivel central. En el establecimiento se registrará el Formulario ITS 1 y, a partir de esos datos, el sistema generará automáticamente el consolidado ITS 2. Los niveles superiores revisarán únicamente información consolidada, podrán devolver reportes con observaciones o aprobarlos según sus competencias.",
)
add_body(
    doc,
    "La plataforma incorporará administración de usuarios y territorios, validaciones de calidad, versionamiento de reportes, exportación a formatos oficiales, tableros de indicadores y mapas de producción y procedencia. Cada usuario visualizará solamente las funciones y los datos correspondientes a su rol y ámbito territorial.",
)
add_callout(
    doc,
    "Principio de privacidad",
    "El dato individual del Formulario ITS 1 permanecerá en el establecimiento que lo genera. Los niveles municipal, regional y central consultarán únicamente información consolidada o agregada.",
)

doc.add_heading("4. Funcionalidades principales", level=1)
for item in (
    "Inicio de sesión y control de acceso por rol, institución y alcance territorial.",
    "Administración de regiones, municipios, establecimientos de salud, usuarios, permisos y catálogos.",
    "Captura, consulta y corrección del Formulario ITS 1 dentro del establecimiento autorizado.",
    "Generación automática del Formulario ITS 2 a partir de los registros capturados.",
    "Registro del total de atenciones por grupos de edad para el cálculo de tasas e indicadores.",
    "Bandejas de revisión con estados de borrador, enviado, en revisión, devuelto, aprobado y cerrado.",
    "Consolidación automática de reportes en los niveles municipal, regional y nacional.",
    "Generación de reportes mensuales, trimestrales, semestrales, anuales y comparativos.",
    "Exportación de reportes oficiales en Excel y PDF.",
    "Tableros con indicadores, gráficos, filtros por período y semana epidemiológica.",
    "Mapas de establecimientos, procedencia, captación, cobertura y pertenencia al Área Geográfica de Influencia (AGI).",
    "Auditoría de acciones sensibles y conservación de versiones de los reportes.",
):
    add_bullet(doc, item)

doc.add_page_break()

doc.add_heading("5. Alcance del proyecto", level=1)
doc.add_heading("5.1 Alcance de la primera versión", level=2)
add_body(
    doc,
    "La primera versión se desarrollará y validará como piloto en el municipio de Puerto Cortés. Incluirá la coordinación municipal y 12 establecimientos de salud, permitiendo comprobar el flujo completo desde la captura del ITS 1 hasta la revisión y consolidación municipal. La arquitectura y el modelo de datos quedarán preparados para incorporar posteriormente la Región Sanitaria de Cortés y el nivel central.",
)
add_body(doc, "El piloto comprenderá:")
for item in (
    "Configuración de la estructura territorial y de los usuarios participantes.",
    "Implementación de la captura ITS 1 y generación ITS 2.",
    "Flujo de envío, revisión, devolución, corrección y aprobación municipal.",
    "Reportes mensuales y anuales esenciales, exportación en Excel/PDF y tablero básico de indicadores.",
    "Mapa municipal con los establecimientos activos y visualización agregada de captación y procedencia.",
    "Registro de auditoría, reglas de privacidad y validaciones de calidad.",
):
    add_bullet(doc, item)

doc.add_heading("5.2 Límites de la primera versión", level=2)
add_body(
    doc,
    "Para mantener un alcance viable dentro del período académico, la primera versión no contempla un despliegue operativo en todo el país, una aplicación móvil nativa ni la sustitución del expediente clínico institucional. Las integraciones con otros sistemas de salud, la autenticación institucional centralizada y las funciones geográficas avanzadas quedarán sujetas a la disponibilidad de servicios, catálogos oficiales y autorizaciones de la institución.",
)

doc.add_heading("6. Tecnologías que se utilizarán", level=1)
tech_table = doc.add_table(rows=1, cols=3)
set_table_geometry(tech_table, [2300, 2700, 4360])
tech_table.style = "Table Grid"
headers = ("Componente", "Tecnología", "Uso previsto")
for idx, value in enumerate(headers):
    cell = tech_table.rows[0].cells[idx]
    shade_cell(cell, LIGHT)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(value)
    set_font(r, size=9.5, color=INK, bold=True)
set_repeat_table_header(tech_table.rows[0])
rows = [
    ("Interfaz web", "Angular y TypeScript", "Pantallas, formularios, tableros y experiencia de usuario."),
    ("Servidor", "NestJS sobre Node.js", "Reglas de negocio, validaciones, permisos, consolidación y auditoría."),
    ("Base de datos", "PostgreSQL en Supabase", "Persistencia segura de usuarios, catálogos, registros y reportes."),
    ("Datos geográficos", "PostGIS y Leaflet", "Coordenadas, territorios, establecimientos y mapas interactivos."),
    ("Gráficos", "Apache ECharts", "Indicadores y visualizaciones interactivas."),
    ("Documentos", "ExcelJS y LibreOffice", "Generación de archivos Excel y conversión controlada a PDF."),
    ("Desarrollo", "Git y GitHub", "Control de versiones, respaldo del código y seguimiento de cambios."),
]
for values in rows:
    cells = tech_table.add_row().cells
    for idx, value in enumerate(values):
        cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(value)
        set_font(r, size=9.3)
set_table_geometry(tech_table, [2300, 2700, 4360])

add_body(
    doc,
    "Como componentes de apoyo se evaluarán Prisma para acceso tipado a datos, Redis y BullMQ para procesar importaciones o exportaciones pesadas, Docker para estandarizar los entornos y herramientas de observabilidad para registrar errores y rendimiento. Su incorporación dependerá de las necesidades comprobadas durante el piloto.",
)

doc.add_heading("7. Justificación del proyecto", level=1)
add_body(
    doc,
    "SIGVITS es una solución útil porque responde a una necesidad operativa concreta: transformar un proceso fragmentado y manual en un flujo institucional único, trazable y seguro. La captura se realizará una sola vez en el establecimiento y el resto de los niveles trabajará con consolidados generados por el sistema. Esto permitirá disminuir errores de transcripción, acortar los tiempos de preparación de informes y mejorar la oportunidad de la información.",
)
add_body(
    doc,
    "El proyecto también aporta valor a la toma de decisiones. Los indicadores, comparativos y mapas facilitarán reconocer variaciones por período, concentración de casos, procedencias externas y diferencias entre producción del establecimiento y cobertura territorial. La auditoría y el versionamiento fortalecerán la rendición de cuentas y permitirán conocer el estado real de cada reporte.",
)
add_body(
    doc,
    "La propuesta es viable porque utiliza tecnologías web maduras, ampliamente documentadas y compatibles con un desarrollo incremental. El piloto en Puerto Cortés limita inicialmente la cantidad de usuarios, establecimientos y datos que deben configurarse, por lo que permite validar las reglas institucionales sin asumir desde el inicio el costo y la complejidad de una implementación nacional. Una vez evaluados los resultados, la misma arquitectura podrá ampliarse progresivamente a otros municipios y regiones.",
)
add_callout(
    doc,
    "Resultado esperado",
    "Un prototipo funcional validado con el flujo de Puerto Cortés, capaz de demostrar la captura, consolidación, revisión, protección, análisis y exportación de información del programa ITS.",
)

doc.add_heading("8. Conclusión", level=1)
add_body(
    doc,
    "El desarrollo de SIGVITS permitirá demostrar cómo una plataforma web puede mejorar la gestión del programa ITS sin alterar la jerarquía institucional ni exponer información individual en niveles no autorizados. Su enfoque por etapas, su delimitación territorial y el uso de tecnologías conocidas hacen que la propuesta sea pertinente para una práctica profesional y, al mismo tiempo, tenga potencial de continuidad institucional.",
)

doc.core_properties.title = "Propuesta de Proyecto SIGVITS"
doc.core_properties.subject = "Sistema de Gestión y Vigilancia de Infecciones de Transmisión Sexual"
doc.core_properties.author = "Estudiante de práctica profesional"
doc.core_properties.keywords = "SIGVITS, ITS, salud pública, Puerto Cortés, propuesta académica"

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
